import os
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class Act:
    def __init__(self, client, name_technique, sn_technique):
        self.name_docx = "Акт-Расписка.docx"

        self.you_list_employee = [
            ['login', 'full name', 'pos'],
        ]

        self.list_month = {
            '01' : 'января', '02' : 'февраля','03' : 'марта',
            '04' : 'апреля', '05' : 'мая', '06' : 'июня',
            '08' : 'августа', '09' : 'сентября', '07' : 'июля',
            '10' : 'октября', '11' : 'ноября', '12' : 'декабря',
        }

        self.date_param = []
        self.worker = client
        self.it_user = []

        self.number = []
        self.technic = name_technique
        self.sn = sn_technique

        if len(self.technic) == len(self.sn):
            self.count_technic = len(self.technic)

        self.units = []
        self.quantity = []

        self.other = []


    def central_command(self):
        self.current_date()
        self.data_it()
        self.data_processing()
        self.completion()


    def current_date(self):
        day = datetime.now().strftime('%d')
        month = datetime.now().strftime('%m')
        month_name = self.list_month[month]
        year = datetime.now().strftime('%y')
        self.date_param.append(day)
        self.date_param.append(month_name)
        self.date_param.append(year)


    def data_it(self):
        current_user = os.getlogin()
        for employee in self.you_list_employee:
            if employee[0] == current_user or employee[1] == current_user:
                self.it_user.extend(employee[2:])
        self.it_user.reverse()


    def data_processing(self):
        for i in range(self.count_technic):
            self.number.append(i + 1)

        for i in range(self.count_technic):
            self.units.append('шт.')

        for i in range(self.count_technic):
            self.quantity.append(1)

        for i in range(len(self.number)):
            item = [
                self.number[i],
                self.technic[i],
                self.sn[i],
                self.units[i],
                self.quantity[i]
            ]

            self.other.append(item)


    def completion(self):
        doc = Document(self.name_docx)

        self.act_update(doc, 1, [0, 0], [2, 6], self.it_user, 'null')
        self.act_update(doc, 2, [0, 0], [2, 6], self.worker, 'null')
        self.act_update(doc, 3, [0, 0, 0], [2, 4, 6], self.date_param, 'null')

        self.act_update(doc, 0, [], [], self.technic, self.sn)

        date = datetime.now().strftime('%d.%m.%Y')
        doc.save(f"{self.worker[1]} {date}.docx")


    def act_update(self, doc, table_number, row, col, data1, data2):
        tab = doc.tables[table_number]

        if table_number == 0:
            if self.count_technic + 2 >= len(tab.rows):
                rows_to_add = self.count_technic + 2 - len(tab.rows)
                for _ in range(rows_to_add):
                    tab.add_row()

        if data2 == 'null':
            for i in range(len(row)):
                cell = tab.cell(row[i], col[i])
                cell.text = ''
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(str(data1[i]))
                run.font.size = Pt(12)
                run.font.name = 'XO Thames'
        else:
            col = [0, 1, 2, 3, 4]
            row = []
            for i in range(self.count_technic):
                row.append(2 + i)

            for i in range(len(self.number)):
                for j in range(len(col)):
                    cell = tab.cell(row[i], col[j])
                    cell.text = ''
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(str(self.other[i][j]))
                    run.font.size = Pt(12)
                    run.font.name = 'XO Thames'
